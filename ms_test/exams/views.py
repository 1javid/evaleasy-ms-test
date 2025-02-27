import random
import zipfile
from io import BytesIO
import os
from django.conf import settings

from django.http import JsonResponse, HttpResponse, Http404, FileResponse
from django.db import transaction
from django.shortcuts import get_object_or_404

from rest_framework import generics, status
from rest_framework.response import Response
from rest_framework.views import APIView

from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENTATION


from .models import *
from .serializers import *

from .models import Subject, QuestionPool, Question
from .serializers import SubjectSerializer, QuestionPoolSerializer, QuestionSerializer

# Endpoint: Create Subject
class CreateSubjectView(generics.CreateAPIView):
    queryset = Subject.objects.all()
    serializer_class = SubjectSerializer

class ListSubjectsView(generics.ListAPIView):
    serializer_class = SubjectSerializer

    def get_queryset(self):
        created_by = self.request.query_params.get('created_by')
        if created_by:
            return Subject.objects.filter(created_by=created_by)
        return Subject.objects.all()

class ListTestQuestionsByAssessmentIdView(generics.ListAPIView):
    serializer_class = QuestionSerializer

    def get_queryset(self):
        assessment_id = self.kwargs['assessment_id']
        test_questions = TestQuestion.objects.filter(assessment_id=assessment_id)
        question_ids = test_questions.values_list('question_id', flat=True)
        return Question.objects.filter(id__in=question_ids)

class RetrieveSubjectView(generics.RetrieveAPIView):
    queryset = Subject.objects.all()
    serializer_class = SubjectSerializer

# Endpoint: Create Question (with nested Answers)
class CreateQuestionWithAnswersView(generics.CreateAPIView):
    queryset = Question.objects.all()
    serializer_class = QuestionSerializer

class CreateManyQuestionsView(generics.CreateAPIView):
    serializer_class = BulkQuestionSerializer

    def get_serializer_context(self):
        context = super().get_serializer_context()
        context['question_pool'] = self.kwargs['question_pool']
        return context

    def create(self, request, *args, **kwargs):
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        self.perform_create(serializer)
        headers = self.get_success_headers(serializer.data)
        return Response(serializer.data, status=status.HTTP_201_CREATED, headers=headers)

class DeleteQuestionFromPoolView(generics.DestroyAPIView):
    queryset = Question.objects.all()
    serializer_class = QuestionSerializer

    def get_object(self):
        question_pool_id = self.kwargs['question_pool']
        question_id = self.kwargs['id']
        return get_object_or_404(Question, id=question_id, question_pool_id=question_pool_id)

class ListQuestionsByQuestionPoolView(generics.ListAPIView):
    serializer_class = QuestionSerializer

    def get_queryset(self):
        question_pool_id = self.kwargs['question_pool_id']
        return Question.objects.filter(question_pool_id=question_pool_id)

# Endpoint: Create Question Pool
class CreateQuestionPoolView(generics.CreateAPIView):
    queryset = QuestionPool.objects.all()
    serializer_class = QuestionPoolSerializer

class ListQuestionPoolsBySubjectView(generics.ListAPIView):
    serializer_class = QuestionPoolSerializer

    def get_queryset(self):
        subject_id = self.kwargs['subject_id']
        return QuestionPool.objects.filter(subject_id=subject_id)
    
class RetrieveTestByAssessmentIdView(generics.RetrieveAPIView):
    queryset = Test.objects.all()
    serializer_class = TestSerializer
    lookup_field = 'assessment_id'

class ListTestsBySubjectView(generics.ListAPIView):
    serializer_class = TestSerializer

    def get_queryset(self):
        subject_id = self.kwargs['subject_id']
        return Test.objects.filter(subject_id=subject_id)
    
class ListGroupIdsBySubjectView(generics.ListAPIView):
    serializer_class = GroupIdSerializer

    def get_queryset(self):
        subject_id = self.kwargs['subject_id']
        return Test.objects.filter(subject_id=subject_id).values('group_id').distinct()
    
class ListTestsByGroupIdView(generics.ListAPIView):
    serializer_class = TestSerializer

    def get_queryset(self):
        group_id = self.kwargs['group_id']
        return Test.objects.filter(group_id=group_id)

class GetDownloadLinkByGroupIdView(APIView):
    def get(self, request, *args, **kwargs):
        group_id = self.kwargs.get('group_id')
        
        # Retrieve all generated links associated with the group_id
        generated_links = GeneratedTestLink.objects.filter(test__group_id=group_id)
        if not generated_links.exists():
            return Response(
                {"detail": "No tests found for the given group ID."},
                status=404
            )
        
        # Create a ZIP file in memory
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w') as zip_file:
            for generated_link in generated_links:
                test = generated_link.test
                file_name = f"{test.name}_Variant_{test.variant}.docx"
                
                # If exam_file is a FileField, read its content
                if hasattr(generated_link.exam_file, 'read'):
                    file_content = generated_link.exam_file.read()
                else:
                    file_content = generated_link.exam_file
                
                zip_file.writestr(file_name, file_content)
        
        # Reset buffer pointer to the beginning
        zip_buffer.seek(0)
        
        # Create the HTTP response with the zip file
        response = HttpResponse(zip_buffer.getvalue(), content_type='application/zip')
        response['Content-Disposition'] = f'attachment; filename="{group_id}_tests.zip"'
        return response

def generate_unique_id():
    """Generate a unique 5-digit id as a string."""
    while True:
        unique_id = str(random.randint(10000, 99999))
        if not Test.objects.filter(group_id=unique_id).exists() and not Test.objects.filter(assessment_id=unique_id).exists():
            return unique_id

def generate_unique_group_id():
    """Generate a unique 5-digit group id as a string."""
    return generate_unique_id()

def generate_unique_assessment_id():
    """Generate a unique 5-digit assessment id as a string."""
    return generate_unique_id()

def create_word_file(subject_name, assessment_id, test_name, variant, sorted_test_questions, answer_sheet_image_path):
    """
    Create a Word document with the given header and test questions.
    sorted_test_questions: list of tuples (position, question) in order.
    Each question's answers are labeled from A to E (up to 5 answers).
    The question text will also show the point value at the end.
    """
    document = Document()

    # Header
    document.add_heading(f"{subject_name} - {test_name}", level=0)
    document.add_paragraph(f"Assessment ID: {assessment_id}")
    document.add_paragraph(f"Variant: {variant}")
    document.add_paragraph("")  # empty line

    # For each test question, add the question text and its answers.
    for pos, question in sorted_test_questions:
        # Append the point cost to the question text.
        # For example: "1. What is 2 + 2? (1 pt.)"
        question_text_with_points = f"{pos}. {question.text} ({question.default_score} pt.)"
        document.add_paragraph(question_text_with_points)
        
        # Retrieve answers and label them A, B, C, etc.
        answers = list(question.answers.all())
        letters = ['A', 'B', 'C', 'D', 'E']
        for idx, answer in enumerate(answers[:5]):
            document.add_paragraph(f"   ({letters[idx]}) {answer.text}")

    # Add a new section for the answer sheet image
    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENTATION.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)

    # Add the answer sheet image
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(answer_sheet_image_path, width=Inches(8.5), height=Inches(11))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save document to a BytesIO stream and return the bytes
    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream.read()

def download_word_file(request, test_id):
    """
    Retrieve the generated Word file for a given test and serve it as a downloadable file.
    """
    try:
        # Assuming there is only one GeneratedTestLink per test.
        generated_link = GeneratedTestLink.objects.get(test__id=test_id)
    except GeneratedTestLink.DoesNotExist:
        raise Http404("Word file not found for this test.")

    # Prepare the HTTP response with the binary file.
    response = HttpResponse(
        generated_link.exam_file,
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="test_{test_id}.docx"'
    return response

class GenerateTestView(APIView):
    """
    This endpoint receives instructions to generate a test.
    The instructor provides the subject, instructor_id, test name,
    a list of variants, and a list of question selections.
    Each question selection is a dict with a question_pool id and a list of positions.
    For the first variant the questions will be assigned per the given positions;
    for additional variants the questions from each pool are shuffled.
    A Word file is generated and stored.
    """
    serializer_class = TestGenerationSerializer

    @transaction.atomic
    def post(self, request, *args, **kwargs):
        serializer = TestGenerationSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        data = serializer.validated_data

        # Retrieve the subject (assumes subject exists)
        subject = get_object_or_404(Subject, id=data['subject'])
        instructor_id = data['instructor_id']
        test_name = data['name']
        variants = data['variants']
        question_selections = data['question_selections']
        notes = data.get('notes')
        instructions = data.get('instructions')

        results = []

        # Generate a unique group ID once for all variants
        group_id = generate_unique_group_id()

        # For each variant, build a test
        for variant in variants:
            # Generate a unique assessment ID for each variant
            assessment_id = generate_unique_assessment_id()

            # Dictionary mapping question position -> Question instance.
            test_questions_mapping = {}

            # For each question selection, pick questions from the indicated pool.
            for qs in question_selections:
                qp_id = qs['question_pool']
                positions = qs['positions']  # list of desired positions
                # Retrieve the QuestionPool instance
                question_pool = get_object_or_404(QuestionPool, id=qp_id)
                available_questions = list(question_pool.questions.all())
                if len(available_questions) < len(positions):
                    return Response(
                        {"detail": f"Not enough questions in QuestionPool {qp_id} to fill positions {positions}"},
                        status=status.HTTP_400_BAD_REQUEST,
                    )
                # Randomly sample distinct questions for the count required
                selected_questions = random.sample(available_questions, len(positions))
                # For the base variant, assign in the order provided.
                # For additional variants, shuffle the selected questions.
                if variant != variants[0]:
                    random.shuffle(selected_questions)
                # Assign each question to its corresponding position.
                for pos, question in zip(positions, selected_questions):
                    if pos in test_questions_mapping:
                        return Response(
                            {"detail": f"Duplicate position {pos} specified across question selections."},
                            status=status.HTTP_400_BAD_REQUEST,
                        )
                    test_questions_mapping[pos] = question

            # Ensure questions are sorted by their position
            sorted_test_questions = sorted(test_questions_mapping.items(), key=lambda x: x[0])

            # Create a Test record with the generated group ID and assessment ID.
            test_obj = Test.objects.create(
                subject=subject,
                instructor_id=instructor_id,
                group_id=group_id,
                assessment_id=assessment_id,
                name=test_name,
                variant=variant,
                notes=notes,
                instructions=instructions,
            )

            # Create TestQuestion records
            for pos, question in sorted_test_questions:
                TestQuestion.objects.create(
                    test=test_obj,
                    question=question,
                    position=pos,
                    assessment_id=assessment_id  # Add this line
                )

            # Generate the Word file for this test
            word_file_bytes = create_word_file(
                subject_name=subject.name,
                assessment_id=assessment_id,
                test_name=test_name,
                variant=variant,
                sorted_test_questions=sorted_test_questions,
                answer_sheet_image_path="./exams/img/blank_sheet.jpg"
            )

            # Save the generated Word file
            GeneratedTestLink.objects.create(
                test=test_obj,
                exam_file=word_file_bytes,
                assessment_id=assessment_id
            )

            results.append({
                "test_id": test_obj.id,
                "group_id": group_id,
                "assessment_id": assessment_id,
                "variant": variant,
            })

        return Response({"generated_tests": results}, status=status.HTTP_201_CREATED)

class RegenerateTestFileView(APIView):
    """
    Regenerates the Word file for an already created test.
    The test is identified by test_id.
    The new file is generated using the existing TestQuestion records.
    """
    @transaction.atomic
    def post(self, request, test_id, *args, **kwargs):
        # Retrieve the test record.
        test_obj = get_object_or_404(Test, id=test_id)
        
        # Retrieve associated test questions, sorted by their position.
        test_question_qs = TestQuestion.objects.filter(test=test_obj).order_by('position')
        if not test_question_qs.exists():
            return Response({"detail": "No test questions found for this test."},
                            status=status.HTTP_400_BAD_REQUEST)
        
        # Build a list of (position, question) tuples.
        sorted_test_questions = [(tq.position, tq.question) for tq in test_question_qs]
        
        # Generate a new Word file using the same helper function.
        word_file_bytes = create_word_file(
            subject_name=test_obj.subject.name,
            assessment_id=test_obj.assessment_id,
            test_name=test_obj.name,
            variant=test_obj.variant,
            sorted_test_questions=sorted_test_questions,
        )
        
        # Update the existing GeneratedTestLink record, or create one if it doesn't exist.
        generated_link = test_obj.generated_links.first()
        if generated_link:
            generated_link.exam_file = word_file_bytes
            generated_link.save()
        else:
            GeneratedTestLink.objects.create(test=test_obj, exam_file=word_file_bytes)
        
        return Response({"detail": "Test regenerated successfully."}, status=status.HTTP_200_OK)
    
def correct_answers_view(request, assessment_id):
    """
        Endpoint: GET /api/tests/<assessment_id>/correct_answers/
        
        Returns a JSON response with:
        - "correct_answers": mapping of question positions to the correct answer letter (A-E).
        - "points": mapping of question positions to the question's default score.
        
        Example response:
        {
            "correct_answers": {
                "1": "B",
                "2": "C",
                "3": "B",
                // ...
            },
            "points": {
                "1": 1.0,
                "2": 1.0,
                "3": 2.0,
                // ...
            }
        }
    """
    # Retrieve the Test object using its unique assessment_id.
    test_obj = get_object_or_404(Test, assessment_id=assessment_id)
    
    # Retrieve TestQuestion records for this test, ordered by the position field.
    test_questions = test_obj.test_questions.all().order_by('position')
    
    correct_answers_mapping = {}
    points_mapping = {}
    
    # Define letters to label answers (up to 5 possible answers)
    letters = ['A', 'B', 'C', 'D', 'E']
    
    for tq in test_questions:
        position_key = str(tq.position)  # Convert position to string for the JSON mapping keys.
        question = tq.question
        
        # Retrieve answers ordered by id (as ensured by Answer.Meta.ordering)
        answers_list = list(question.answers.all())
        correct_letter = None
        
        # Determine which answer is marked as correct.
        for idx, answer in enumerate(answers_list):
            if answer.is_correct:
                if idx < len(letters):
                    correct_letter = letters[idx]
                else:
                    correct_letter = "?"  # Fallback if more than 5 answers exist.
                break
        
        # If no correct answer is found, mark as "N/A".
        if correct_letter is None:
            correct_letter = "N/A"
        
        correct_answers_mapping[position_key] = correct_letter
        points_mapping[position_key] = float(question.default_score)
    
    response_data = {
        "correct_answers": correct_answers_mapping,
        "points": points_mapping,
    }
    
    return JsonResponse(response_data)
